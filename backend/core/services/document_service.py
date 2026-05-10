"""
Document processing service.

Handles file validation, text extraction from PDF, DOCX, TXT, Excel, and image files,
and text normalization for downstream processing.

Merges Banking's robust extraction (tables + OCR confidence) with Compliance's
file support (.txt, .md).
"""

import logging
import os
import re
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

from core.config import settings

logger = logging.getLogger(__name__)

# Try to import pytesseract; it requires Tesseract-OCR to be installed
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract not available. OCR features will be disabled.")


class DocumentService:
    """
    Service for document file handling and text extraction.

    Supports PDF, DOCX, TXT, Excel, and image file formats with proper
    validation and error handling.
    """

    SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".xlsx", ".xls"}

    def __init__(self, upload_dir: str = None) -> None:
        """Initialize the document service."""
        self.upload_dir = upload_dir or settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)

    def validate_file_type(self, filename: str) -> bool:
        """
        Validate that the file extension is supported.
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported types: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )
        return True

    def validate_file_size(self, file_size: int, max_size_mb: int = 10) -> bool:
        """
        Validate that the file size is within limits.
        """
        max_bytes = max_size_mb * 1024 * 1024
        if file_size > max_bytes:
            raise ValueError(f"File size too large. Maximum size is {max_size_mb}MB.")
        return True

    def extract_text_and_tables(self, file_path: str) -> tuple[str, str, float]:
        """
        Extract both raw text and structured table data from a document.

        Args:
            file_path: Path to the document.

        Returns:
            Tuple of (plain_text, table_text, ocr_confidence).
            ocr_confidence is 1.0 for non-image files, 0.0 if OCR skipped.
        """
        if not file_path or not isinstance(file_path, str):
            logger.error("Invalid file path provided")
            return "", "", 0.0
        
        if not os.path.exists(file_path):
            logger.error("File not found: %s", file_path)
            return "", "", 0.0
            
        ext = Path(file_path).suffix.lower()
        ocr_confidence = 1.0
        table_text = ""
        plain_text = ""

        try:
            if ext == ".pdf":
                plain_text = self.extract_text_from_pdf(file_path)
                try:
                    table_text = self.extract_tables_from_pdf(file_path)
                except Exception as table_err:
                    logger.debug("Table extraction failed (non-critical): %s", table_err)
            elif ext == ".docx":
                plain_text = self.extract_text_from_docx(file_path)
            elif ext in {".png", ".jpg", ".jpeg"}:
                plain_text, ocr_confidence = self.extract_text_from_image(file_path)
            elif ext in {".xlsx", ".xls"}:
                plain_text = self.extract_from_xlsx(file_path)
            elif ext in {".txt", ".md"}:
                plain_text = self.extract_text_from_txt(file_path)
            else:
                logger.error("Unsupported file extension: %s", ext)
                return "", "", 0.0
            
            if not plain_text or not isinstance(plain_text, str):
                logger.warning("Failed to extract text from %s", file_path)
                return "", "", 0.0
        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, str(e))
            return "", "", 0.0

        return self.normalize_text(plain_text), self.normalize_text(table_text), ocr_confidence

    def extract_text(self, file_path: str) -> str:
        """Convenience method that returns only the combined plain text."""
        plain_text, table_text, _ = self.extract_text_and_tables(file_path)
        combined = plain_text
        if table_text:
            combined += "\n\n" + table_text
        return combined.strip()

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from a plain TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="iso-8859-1") as f:
                return f.read()

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using pdfplumber."""
        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.debug("No text extracted from page %d of %s", page_num, file_path)
        return "\n".join(text_parts)

    def extract_tables_from_pdf(self, file_path: str) -> str:
        """Extract tabular data from a PDF file using pdfplumber."""
        table_parts: list[str] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    if not tables:
                        continue
                    table_parts.append(f"--- Page {page_num} Tables ---")
                    for tbl_idx, table in enumerate(tables, start=1):
                        if not table:
                            continue
                        table_parts.append(f"Table {tbl_idx}:")
                        for row in table:
                            formatted = " | ".join(
                                (str(cell).strip() if cell else "") for cell in row
                            )
                            if formatted.strip("| "):
                                table_parts.append(formatted)
                        table_parts.append("")
        except Exception as e:
            logger.warning("Table extraction failed for %s: %s", file_path, str(e))
        return "\n".join(table_parts)

    def extract_from_xlsx(self, file_path: str) -> str:
        """Extract text content from an Excel XLSX/XLS file."""
        try:
            import openpyxl  # noqa: PLC0415
        except ImportError:
            raise RuntimeError(
                "openpyxl is required for Excel support. Run: pip install openpyxl"
            )

        parts: list[str] = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    formatted = " | ".join(
                        (str(cell).strip() if cell is not None else "") for cell in row
                    )
                    if formatted.strip("| "):
                        parts.append(formatted)
                        row_count += 1
                    if row_count >= 2000:
                        parts.append("[... truncated ...]")
                        break
                parts.append("")
            wb.close()
        except Exception as e:
            raise RuntimeError(f"Failed to read Excel file: {str(e)}") from e
        return "\n".join(parts)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file using python-docx."""
        doc = DocxDocument(file_path)
        text_parts: list[str] = []

        def _append_paragraphs(paragraphs) -> None:
            for p in paragraphs:
                if getattr(p, "text", "").strip():
                    text_parts.append(p.text)

        def _append_tables(tables) -> None:
            for table in tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

        _append_paragraphs(doc.paragraphs)
        _append_tables(doc.tables)

        try:
            for section in doc.sections:
                header = getattr(section, "header", None)
                if header:
                    _append_paragraphs(header.paragraphs)
                    _append_tables(header.tables)
                footer = getattr(section, "footer", None)
                if footer:
                    _append_paragraphs(footer.paragraphs)
                    _append_tables(footer.tables)
        except Exception as e:
            logger.debug("DOCX header/footer extraction failed for %s: %s", file_path, e)

        return "\n".join(text_parts)

    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Apply image preprocessing to improve OCR accuracy."""
        image = image.convert("L")
        image = ImageOps.autocontrast(image, cutoff=2)
        image = ImageEnhance.Contrast(image).enhance(1.5)
        image = image.filter(ImageFilter.SHARPEN)
        return image

    def extract_text_from_image(self, file_path: str) -> tuple[str, float]:
        """Extract text from an image file using OCR (Tesseract) with preprocessing."""
        if not TESSERACT_AVAILABLE:
            raise RuntimeError(
                "OCR is not available. Install Tesseract-OCR and pytesseract to process images."
            )

        image = Image.open(file_path)
        image = self._preprocess_image_for_ocr(image)

        try:
            import pandas as pd  # noqa: PLC0415
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DATAFRAME)
            valid = data[(data["conf"] > 0) & (data["text"].str.strip().ne(""))]
            avg_conf = float(valid["conf"].mean()) / 100.0 if not valid.empty else 0.0
            text = " ".join(valid["text"].astype(str).tolist())
        except Exception:
            text = pytesseract.image_to_string(image)
            avg_conf = 0.75

        return text, avg_conf

    def normalize_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\S\n]+", " ", text)
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()

    async def save_upload(self, filename: str, content: bytes) -> str:
        """Save uploaded file content to disk."""
        import uuid
        safe_name = f"{uuid.uuid4().hex}_{filename}"
        file_path = os.path.join(self.upload_dir, safe_name)

        with open(file_path, "wb") as f:
            f.write(content)

        logger.info("Saved upload: %s -> %s", filename, file_path)
        return file_path

    def cleanup_file(self, file_path: str) -> None:
        """Remove a temporary file from disk."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug("Cleaned up file: %s", file_path)
        except OSError as e:
            logger.warning("Failed to clean up file %s: %s", file_path, str(e))
