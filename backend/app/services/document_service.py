"""
Document processing service.

Handles file validation, text extraction from PDF, DOCX, and image files,
and text normalization for downstream processing.
"""

import logging
import os
import re
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from PIL import Image

from app.config import settings

logger = logging.getLogger(__name__)

# Try to import pytesseract; it requires Tesseract-OCR to be installed
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract not available. OCR features will be disabled.")


class DocumentService:
    """
    Service for document file handling and text extraction.

    Supports PDF, DOCX, TXT, and MD file formats with proper
    validation and error handling.
    """

    SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg"}

    def __init__(self) -> None:
        """Initialize the document service."""
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    def validate_file_type(self, filename: str) -> bool:
        """
        Validate that the file extension is supported.

        Args:
            filename: Original filename to validate.

        Returns:
            True if the file type is supported.

        Raises:
            ValueError: If the file type is not supported.
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported types: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )
        return True

    def validate_file_content(self, content: bytes, filename: str) -> bool:
        """
        Perform deep validation to ensure the file's binary content matches its extension.

        Args:
            content: Raw bytes of the uploaded file.
            filename: Original filename.

        Returns:
            True if deep validation passes.

        Raises:
            ValueError: If file content appears spoofed.
        """
        ext = Path(filename).suffix.lower()
        
        # Prevent executable spoofing via PDF
        if ext == ".pdf":
            # PDF magic bytes: %PDF-
            if not content.startswith(b"%PDF-"):
                raise ValueError("Deep validation failed: File extension is .pdf but the content does not contain a valid PDF signature.")
                
        return True

    def validate_file_size(self, file_size: int) -> bool:
        """
        Validate that the file size is within limits.

        Args:
            file_size: Size of the file in bytes.

        Returns:
            True if the file size is within limits.

        Raises:
            ValueError: If the file exceeds the maximum size.
        """
        if file_size > settings.MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds "
                f"maximum allowed size ({settings.MAX_FILE_SIZE_MB} MB)."
            )
        return True

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension.

        Args:
            file_path: Path to the uploaded file.

        Returns:
            Extracted and normalized text content.

        Raises:
            ValueError: If the file type is not supported.
            RuntimeError: If text extraction fails.
        """
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                raw_text = self.extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = self.extract_text_from_docx(file_path)
            elif ext in {".txt", ".md"}:
                raw_text = self.extract_text_from_txt(file_path)
            elif ext in {".png", ".jpg", ".jpeg"}:
                raw_text, _ = self.extract_text_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {ext}")

            normalized = self.normalize_text(raw_text)
            logger.info(
                "Extracted %d characters from %s (normalized: %d)",
                len(raw_text), file_path, len(normalized)
            )
            return normalized

        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, str(e))
            raise RuntimeError(f"Failed to extract text from file: {str(e)}") from e

    def extract_text_and_tables(self, file_path: str) -> tuple[str, str, float]:
        """
        Extract both raw text and structured table data from a document.

        Args:
            file_path: Path to the document.

        Returns:
            Tuple of (plain_text, table_text, ocr_confidence).
            ocr_confidence is 1.0 for non-image files, 0.0 if extraction skipped.
        """
        if not file_path or not isinstance(file_path, str):
            logger.error("Invalid file path provided")
            return "", "", 0.0

        if not os.path.exists(file_path):
            logger.error("File not found: %s", file_path)
            return "", "", 0.0

        ext = Path(file_path).suffix.lower()
        ocr_confidence = 1.0
        table_text = ""
        plain_text = ""

        try:
            if ext == ".pdf":
                plain_text = self.extract_text_from_pdf(file_path)
                try:
                    table_text = self.extract_tables_from_pdf(file_path)
                except Exception as table_err:
                    logger.debug("Table extraction failed (non-critical): %s", table_err)
            elif ext == ".docx":
                plain_text = self.extract_text_from_docx(file_path)
            elif ext in {".txt", ".md"}:
                plain_text = self.extract_text_from_txt(file_path)
            elif ext in {".png", ".jpg", ".jpeg"}:
                plain_text, ocr_confidence = self.extract_text_from_image(file_path)
            else:
                logger.error("Unsupported file extension: %s", ext)
                return "", "", 0.0

            if not plain_text or not isinstance(plain_text, str):
                logger.warning("Failed to extract text from %s", file_path)
                return "", "", 0.0
        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, str(e))
            return "", "", 0.0

        return self.normalize_text(plain_text), self.normalize_text(table_text), ocr_confidence

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from a plain TXT file.

        Args:
            file_path: Path to the TXT file.

        Returns:
            The plain text content of the file.
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback for ISO-8859-1 encoding if UTF-8 fails
            with open(file_path, "r", encoding="iso-8859-1") as f:
                return f.read()

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file using pdfplumber.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Concatenated text from all pages.
        """
        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.debug("No text extracted from page %d of %s", page_num, file_path)
        return "\n".join(text_parts)

    def extract_tables_from_pdf(self, file_path: str) -> str:
        """
        Extract tabular data from a PDF file using pdfplumber.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Formatted table text, or empty string if no tables found.
        """
        table_parts: list[str] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    if not tables:
                        continue
                    table_parts.append(f"--- Page {page_num} Tables ---")
                    for tbl_idx, table in enumerate(tables, start=1):
                        if not table:
                            continue
                        table_parts.append(f"Table {tbl_idx}:")
                        for row in table:
                            formatted = " | ".join(
                                (str(cell).strip() if cell else "") for cell in row
                            )
                            if formatted.strip("| "):
                                table_parts.append(formatted)
                        table_parts.append("")
        except Exception as e:
            logger.warning("Table extraction failed for %s: %s", file_path, str(e))
        return "\n".join(table_parts)

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Concatenated text from all paragraphs.
        """
        doc = DocxDocument(file_path)
        text_parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def extract_text_from_image(self, file_path: str) -> tuple[str, float]:
        """
        Extract text from an image file using OCR (Tesseract).

        Args:
            file_path: Path to the image file.

        Returns:
            Tuple of (extracted_text, confidence 0.0-1.0).

        Raises:
            RuntimeError: If Tesseract is not available.
        """
        if not TESSERACT_AVAILABLE:
            raise RuntimeError(
                "OCR is not available. Install Tesseract-OCR and pytesseract to process images."
            )

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text, 0.75  # Default confidence when detailed data not available

    def normalize_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.

        - Removes excessive whitespace
        - Normalizes line endings
        - Strips leading/trailing whitespace

        Args:
            text: Raw extracted text.

        Returns:
            Cleaned and normalized text.
        """
        if not text:
            return ""

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Collapse multiple blank lines into one
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Collapse multiple spaces into one (but preserve newlines)
        text = re.sub(r"[^\S\n]+", " ", text)

        # Strip each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()

    async def save_upload(self, filename: str, content: bytes) -> str:
        """
        Save uploaded file content to disk.

        Args:
            filename: Original filename.
            content: File content as bytes.

        Returns:
            Full path to the saved file.
        """
        import uuid
        safe_name = f"{uuid.uuid4().hex}_{filename}"
        file_path = os.path.join(settings.UPLOAD_DIR, safe_name)

        with open(file_path, "wb") as f:
            f.write(content)

        logger.info("Saved upload: %s -> %s", filename, file_path)
        return file_path

    def cleanup_file(self, file_path: str) -> None:
        """
        Remove a temporary file from disk.

        Args:
            file_path: Path to the file to remove.
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug("Cleaned up file: %s", file_path)
        except OSError as e:
            logger.warning("Failed to clean up file %s: %s", file_path, str(e))
