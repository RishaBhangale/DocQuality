"""
Document processing service.

Handles file validation, text extraction from PDF, DOCX, and image files,
and text normalization for downstream processing.
"""

import logging
import os
import re
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

from banking.config import settings

logger = logging.getLogger(__name__)

# Try to import pytesseract; it requires Tesseract-OCR to be installed
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract not available. OCR features will be disabled.")


class DocumentService:
    """
    Service for document file handling and text extraction.

    Supports PDF, DOCX, and image file formats with proper
    validation and error handling.
    """

    SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".xlsx", ".xls"}

    def __init__(self) -> None:
        """Initialize the document service."""
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    def validate_file_type(self, filename: str) -> bool:
        """
        Validate that the file extension is supported.

        Args:
            filename: Original filename to validate.

        Returns:
            True if the file type is supported.

        Raises:
            ValueError: If the file type is not supported.
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported types: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )
        return True

    def validate_file_size(self, file_size: int) -> bool:
        """
        Validate that the file size is within limits.

        Args:
            file_size: Size of the file in bytes.

        Returns:
            True if the file size is within limits.

        Raises:
            ValueError: If the file exceeds the maximum size.
        """
        if file_size > settings.MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds "
                f"maximum allowed size ({settings.MAX_FILE_SIZE_MB} MB)."
            )
        return True

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension.

        Args:
            file_path: Path to the uploaded file.

        Returns:
            Extracted and normalized text content.

        Raises:
            ValueError: If the file type is not supported.
            RuntimeError: If text extraction fails.
        """
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                raw_text = self.extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = self.extract_text_from_docx(file_path)
            elif ext in {".png", ".jpg", ".jpeg"}:
                raw_text, _ = self.extract_text_from_image(file_path)
            elif ext in {".xlsx", ".xls"}:
                raw_text = self.extract_from_xlsx(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {ext}")

            normalized = self.normalize_text(raw_text)
            logger.info(
                "Extracted %d characters from %s (normalized: %d)",
                len(raw_text), file_path, len(normalized)
            )
            return normalized

        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, str(e))
            raise RuntimeError(f"Failed to extract text from file: {str(e)}") from e

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file using pdfplumber.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Concatenated text from all pages.
        """
        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.debug("No text extracted from page %d of %s", page_num, file_path)
        return "\n".join(text_parts)

    def extract_tables_from_pdf(self, file_path: str) -> str:
        """
        Extract tabular data from a PDF file using pdfplumber.

        Tables are rendered as pipe-delimited text rows so the LLM can
        parse structured data (e.g. financial statements, collateral schedules).

        Args:
            file_path: Path to the PDF file.

        Returns:
            Formatted table text, or empty string if no tables found.
        """
        table_parts: list[str] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    if not tables:
                        continue
                    table_parts.append(f"--- Page {page_num} Tables ---")
                    for tbl_idx, table in enumerate(tables, start=1):
                        if not table:
                            continue
                        table_parts.append(f"Table {tbl_idx}:")
                        for row in table:
                            formatted = " | ".join(
                                (str(cell).strip() if cell else "") for cell in row
                            )
                            if formatted.strip("| "):
                                table_parts.append(formatted)
                        table_parts.append("")
        except Exception as e:
            logger.warning("Table extraction failed for %s: %s", file_path, str(e))
        return "\n".join(table_parts)

    def extract_from_xlsx(self, file_path: str) -> str:
        """
        Extract text content from an Excel XLSX/XLS file.

        All sheets are traversed; cell values are rendered as a pipe-delimited
        table so the LLM can reason about structured financial data.

        Args:
            file_path: Path to the Excel file.

        Returns:
            Formatted text representation of all sheets.
        """
        try:
            import openpyxl  # noqa: PLC0415
        except ImportError:
            raise RuntimeError(
                "openpyxl is required for Excel support. Run: pip install openpyxl"
            )

        parts: list[str] = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    formatted = " | ".join(
                        (str(cell).strip() if cell is not None else "") for cell in row
                    )
                    if formatted.strip("| "):
                        parts.append(formatted)
                        row_count += 1
                    if row_count >= 2000:  # Cap at 2000 rows per sheet
                        parts.append("[... truncated ...]")
                        break
                parts.append("")
            wb.close()
        except Exception as e:
            raise RuntimeError(f"Failed to read Excel file: {str(e)}") from e
        return "\n".join(parts)

    def extract_text_and_tables(self, file_path: str) -> tuple[str, str, float]:
        """
        Extract both raw text and structured table data from a document.

        Args:
            file_path: Path to the document.

        Returns:
            Tuple of (plain_text, table_text, ocr_confidence).
            ocr_confidence is 1.0 for non-image files, 0.0 if OCR skipped.
        """
        if not file_path or not isinstance(file_path, str):
            logger.error("Invalid file path provided")
            return "", "", 0.0
        
        if not os.path.exists(file_path):
            logger.error("File not found: %s", file_path)
            return "", "", 0.0
            
        ext = Path(file_path).suffix.lower()
        ocr_confidence = 1.0
        table_text = ""
        plain_text = ""

        try:
            if ext == ".pdf":
                plain_text = self.extract_text_from_pdf(file_path)
                try:
                    table_text = self.extract_tables_from_pdf(file_path)
                except Exception as table_err:
                    logger.debug("Table extraction failed (non-critical): %s", table_err)
            elif ext == ".docx":
                plain_text = self.extract_text_from_docx(file_path)
            elif ext in {".png", ".jpg", ".jpeg"}:
                plain_text, ocr_confidence = self.extract_text_from_image(file_path)
            elif ext in {".xlsx", ".xls"}:
                plain_text = self.extract_from_xlsx(file_path)
            else:
                logger.error("Unsupported file extension: %s", ext)
                return "", "", 0.0
            
            if not plain_text or not isinstance(plain_text, str):
                logger.warning("Failed to extract text from %s", file_path)
                return "", "", 0.0
        except Exception as e:
            logger.error("Text extraction failed for %s: %s", file_path, str(e))
            return "", "", 0.0

        return self.normalize_text(plain_text), self.normalize_text(table_text), ocr_confidence

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Concatenated text from all paragraphs.
        """
        doc = DocxDocument(file_path)
        text_parts: list[str] = []

        def _append_paragraphs(paragraphs) -> None:
            for p in paragraphs:
                if getattr(p, "text", "").strip():
                    text_parts.append(p.text)

        def _append_tables(tables) -> None:
            for table in tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

        # Main document body
        _append_paragraphs(doc.paragraphs)

        # Also extract text from body tables
        _append_tables(doc.tables)

        # Headers/footers often contain version, effective date, approvals, and signatures.
        # Include them to avoid missing critical governance evidence.
        try:
            for section in doc.sections:
                header = getattr(section, "header", None)
                if header:
                    _append_paragraphs(header.paragraphs)
                    _append_tables(header.tables)
                footer = getattr(section, "footer", None)
                if footer:
                    _append_paragraphs(footer.paragraphs)
                    _append_tables(footer.tables)
        except Exception as e:
            logger.debug("DOCX header/footer extraction failed for %s: %s", file_path, e)

        return "\n".join(text_parts)

    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Apply image preprocessing to improve OCR accuracy.

        Steps: convert to greyscale → auto-contrast → contrast enhance → sharpen.
        No external libraries needed beyond Pillow.
        """
        # Greyscale
        image = image.convert("L")
        # Auto-contrast stretches the histogram to fill 0-255
        image = ImageOps.autocontrast(image, cutoff=2)
        # Boost contrast slightly
        image = ImageEnhance.Contrast(image).enhance(1.5)
        # Sharpen edges (helps with degraded scans)
        image = image.filter(ImageFilter.SHARPEN)
        return image

    def extract_text_from_image(self, file_path: str) -> tuple[str, float]:
        """
        Extract text from an image file using OCR (Tesseract) with preprocessing.

        Args:
            file_path: Path to the image file.

        Returns:
            Tuple of (extracted_text, confidence_fraction 0.0-1.0).

        Raises:
            RuntimeError: If Tesseract is not available.
        """
        if not TESSERACT_AVAILABLE:
            raise RuntimeError(
                "OCR is not available. Install Tesseract-OCR and pytesseract to process images."
            )

        image = Image.open(file_path)
        image = self._preprocess_image_for_ocr(image)

        # Use image_to_data to capture per-word confidence scores
        try:
            import pandas as pd  # noqa: PLC0415
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DATAFRAME)
            # Filter words with meaningful confidence (>0) and non-empty text
            valid = data[(data["conf"] > 0) & (data["text"].str.strip().ne(""))]
            avg_conf = float(valid["conf"].mean()) / 100.0 if not valid.empty else 0.0
            text = " ".join(valid["text"].astype(str).tolist())
        except Exception:
            # Fallback: plain image_to_string without confidence
            text = pytesseract.image_to_string(image)
            avg_conf = 0.75  # Assume reasonable quality if data API fails

        return text, avg_conf

    def normalize_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.

        - Removes excessive whitespace
        - Normalizes line endings
        - Strips leading/trailing whitespace

        Args:
            text: Raw extracted text.

        Returns:
            Cleaned and normalized text.
        """
        if not text:
            return ""

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Collapse multiple blank lines into one
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Collapse multiple spaces into one (but preserve newlines)
        text = re.sub(r"[^\S\n]+", " ", text)

        # Strip each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()

    async def save_upload(self, filename: str, content: bytes) -> str:
        """
        Save uploaded file content to disk.

        Args:
            filename: Original filename.
            content: File content as bytes.

        Returns:
            Full path to the saved file.
        """
        import uuid
        safe_name = f"{uuid.uuid4().hex}_{filename}"
        file_path = os.path.join(settings.UPLOAD_DIR, safe_name)

        with open(file_path, "wb") as f:
            f.write(content)

        logger.info("Saved upload: %s -> %s", filename, file_path)
        return file_path

    def cleanup_file(self, file_path: str) -> None:
        """
        Remove a temporary file from disk.

        Args:
            file_path: Path to the file to remove.
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug("Cleaned up file: %s", file_path)
        except OSError as e:
            logger.warning("Failed to clean up file %s: %s", file_path, str(e))
